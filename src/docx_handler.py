# =============================================================================
# src/docx_handler.py — Manipulation des fichiers .docx
#
# Ce module gère toutes les opérations sur les documents Word :
#   - Ouverture en lecture seule (l'original n'est jamais modifié)
#   - Lecture et recherche de paragraphes (y compris dans tableaux et SDT)
#   - Génération du HTML de prévisualisation (document complet ou fenêtre)
#   - Insertion de clauses en mode révision (Track Changes / <w:ins>)
#   - Insertion de clauses en texte brut (sans balisage de révision)
#   - Mise à jour des dates de publication (corps + footer)
#
# Note OOXML — noms de style vs identifiants de style :
#   Word distingue le nom affiché d'un style ("Heading 3", "APU_Heading 3")
#   de son identifiant interne ("Heading3", "APUHeading3"). L'attribut
#   <w:pStyle w:val="..."/> exige l'identifiant, jamais le nom affiché.
#   _resolve_style_id() effectue cette conversion à partir des métadonnées
#   du document (doc.styles), ce qui garantit le bon style quel que soit
#   le template utilisé (classique, APU, RSV…).
#
# Deux fichiers de sortie sont produits par insertion :
#   - _track_changes/<fichier>.docx  : version avec <w:ins> (révision Word)
#   - _texte_brut/<fichier>.docx     : version texte direct, sans révision
#
# Deux types de documents sont supportés :
#   1. Documents "classiques" : paragraphes directement dans <w:body>
#   2. Documents "structurés" : contenu dans des tableaux (<w:td>) ou des
#      contrôles de contenu Word (<w:sdt>/<w:sdtContent>)
#
# python-docx expose uniquement les paragraphes de premier niveau via
# doc.paragraphs. Ce module utilise collect_paragraphs() pour traverser le
# XML complet et capturer tous les paragraphes quelle que soit leur profondeur.
# =============================================================================

import html as _html
import re
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Niveaux d'indentation pour les puces, en twips (1 twip = 1/1440 pouce).
# Niveau 1 = 0.5", niveau 2 = 1", niveau 3 = 1.5"
INDENT_LEVELS: dict[int, int] = {1: 720, 2: 1440, 3: 2160}


# =============================================================================
# Helpers bas niveau sur les éléments <w:p> lxml
#
# Ces fonctions opèrent directement sur des éléments lxml (pas des objets
# python-docx), ce qui permet de traiter des paragraphes à n'importe quelle
# profondeur dans le document XML.
#
# Helpers de style (nécessitent un objet Document python-docx) :
#   _resolve_style_id         : nom affiché → identifiant OOXML (correspondance exacte)
#   _resolve_heading_style_id : idem + fallback par niveau si style absent du document
#   get_para_style_name    : lit l'ID de style d'un <w:p> lxml
#   _is_heading_style      : détecte si un ID de style est un titre
#   get_body_style_near    : cherche le style de corps le plus proche d'un ancre-titre
#   get_para_run_font_size : lit la taille effective des runs (overrides run-level vs style)
#   _find_body_para_near   : comme get_body_style_near mais retourne l'élément <w:p>
#
# Helpers date :
#   _collect_visible_runs : collecte récursive des runs visibles (hors <w:del>)
#   _visible_runs         : API publique de collecte des runs visibles
# =============================================================================

def _para_text(p_elem) -> str:
    """
    Extrait le texte brut d'un élément <w:p> lxml en concaténant tous ses <w:t>.

    Inclut le texte des runs normaux et des insertions Track Changes (<w:ins>),
    mais PAS le texte des suppressions (<w:delText>), ce qui reflète l'état
    "accepté" du document.
    """
    return ''.join(t.text or '' for t in p_elem.iter(qn('w:t')))


# Mots-clés identifiant les styles de titre dans les IDs OOXML.
# Couvre les conventions Word standard ET les variantes maison (APU_Heading, RSV_Heading…).
_HEADING_KEYWORDS = ('heading', 'titre', 'title', 'rsvsection', 'rsvheading')


def get_para_style_name(para_el) -> str:
    """
    Retourne l'ID de style OOXML d'un élément <w:p> lxml, ou '' si absent.

    Note : w:pStyle w:val contient l'ID interne (ex. 'APUHeading3', 'Normal'),
    pas le nom affiché dans Word ('APU_Heading 3'). L'ID est ce qu'on passe
    à <w:pStyle w:val="..."/> lors de l'insertion.
    """
    pPr = para_el.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            return pStyle.get(qn('w:val'), '')
    return ''


def _is_heading_style(style_id: str) -> bool:
    """
    Retourne True si l'ID de style correspond à un titre (toutes familles).

    Exemples détectés : Heading1, Titre2, APUHeading3, RSVHeading1, RSVSection.
    Exemples non détectés : Normal, APUDefault, BodyText, Descriptif.
    """
    sid = style_id.lower()
    return any(k in sid for k in _HEADING_KEYWORDS)


def _resolve_style_id(doc: Document, style_name: str) -> str:
    """
    Convertit un nom de style Word affiché (ex. 'Heading 3', 'APU_Heading 3')
    en son identifiant OOXML (ex. 'Heading3', 'APUHeading3').

    C'est l'ID qui doit figurer dans <w:pStyle w:val="...">, pas le nom affiché.
    En cas d'échec (style inexistant), retourne style_name tel quel.
    """
    for style in doc.styles:
        if style.name == style_name:
            return style.style_id
    return style_name


def _resolve_heading_style_id(doc: Document, style_name: str) -> str:
    """
    Résout le nom d'un style de titre en son identifiant OOXML, avec fallback
    inter-familles de styles.

    Étape 1 — correspondance exacte par nom (comme _resolve_style_id).
    Étape 2 — si le style n'existe pas dans ce document (ex. "Heading 3" dans
               un template AMF qui ne possède que "APU_Heading 3"), cherche un
               style de titre au même niveau numérique dans le document cible.

    Exemples :
        "Heading 3"     dans un doc AMF  → trouve "APU_Heading 3" → "APUHeading3"
        "APU_Heading 3" dans un doc AMF  → trouvé directement      → "APUHeading3"
        "APU_Heading 3" dans un doc classique → trouve "Heading 3" → "Heading3"

    Cela permet de configurer une clause avec "Heading 3" et de l'insérer
    correctement dans les deux familles de documents sans dupliquer les codes.
    """
    # Correspondance exacte par nom affiché
    for style in doc.styles:
        if style.name == style_name:
            return style.style_id

    # Le style est absent de ce document — fallback par niveau numérique
    m = re.search(r'\d+$', style_name.strip())
    if m:
        level = m.group()
        for style in doc.styles:
            sname = style.name.lower()
            if any(k in sname for k in _HEADING_KEYWORDS) and sname.endswith(level):
                return style.style_id

    return style_name  # dernier recours : retourner le nom tel quel


def get_body_style_near(flat_paras: list, anchor_idx: int) -> str:
    """
    Cherche le premier style de corps de texte (non-titre, non vide) au
    voisinage de anchor_idx. Utilisé quand l'ancre est un titre : on hérite
    du style du corps environnant plutôt que du titre lui-même.

    Stratégie : scan en arrière d'abord (le corps précède souvent le titre
    suivant), puis en avant si rien trouvé.

    Returns:
        ID de style (ex. 'APUDefault', 'Normal', 'Descriptif') ou '' si
        aucun paragraphe de corps trouvé dans tout le document.
    """
    # Scan en arrière depuis l'ancre
    for i in range(anchor_idx - 1, -1, -1):
        if not _para_text(flat_paras[i]).strip():
            continue
        style = get_para_style_name(flat_paras[i])
        if style and not _is_heading_style(style):
            return style
    # Scan en avant si rien trouvé
    for i in range(anchor_idx + 1, len(flat_paras)):
        if not _para_text(flat_paras[i]).strip():
            continue
        style = get_para_style_name(flat_paras[i])
        if style and not _is_heading_style(style):
            return style
    return ''


def get_para_run_font_size(para_el) -> int:
    """
    Lit la taille de police effective (en points) du premier run visible non vide.

    En OOXML, la taille peut être définie au niveau du run (w:r/w:rPr/w:sz)
    plutôt qu'au niveau du style. Ce cas est courant quand la mise en forme a
    été appliquée manuellement. Un paragraphe peut donc avoir le style "Normal"
    (10pt) mais afficher du texte en 11pt via des overrides de run.

    Retourne 0 si aucune taille n'est définie explicitement dans les runs
    (le paragraphe hérite alors de la taille définie par son style Word).

    OOXML exprime la taille en demi-points : w:sz val="22" → 11 pt.
    """
    for run_el, text in _visible_runs(para_el):
        if not text.strip():
            continue
        rPr = run_el.find(qn('w:rPr'))
        if rPr is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'), '')
                if val.isdigit():
                    return int(val) // 2  # demi-points → points
    return 0


def _find_body_para_near(flat_paras: list, anchor_idx: int):
    """
    Retourne le premier élément <w:p> de corps de texte (non-titre, non vide)
    au voisinage de anchor_idx.

    Même logique de scan que get_body_style_near, mais retourne l'élément
    <w:p> complet plutôt que son ID de style — permet de lire ses propriétés
    (notamment la taille de police des runs).
    """
    for i in range(anchor_idx - 1, -1, -1):
        if not _para_text(flat_paras[i]).strip():
            continue
        style = get_para_style_name(flat_paras[i])
        if style and not _is_heading_style(style):
            return flat_paras[i]
    for i in range(anchor_idx + 1, len(flat_paras)):
        if not _para_text(flat_paras[i]).strip():
            continue
        style = get_para_style_name(flat_paras[i])
        if style and not _is_heading_style(style):
            return flat_paras[i]
    return None


def _para_html_tag(p_elem) -> str:
    """
    Retourne la balise HTML (h1–h4 ou p) correspondant au style du paragraphe.

    Travaille sur l'ID de style Word (attribut w:val de w:pStyle), pas sur
    le nom complet. Les IDs typiques sont : "Heading1", "Titre2", "1", etc.

    Note : l'ID de style ≠ nom de style. "Heading1" (ID) → "Heading 1" (nom).
    Cette correspondance est suffisante pour la prévisualisation HTML.
    """
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            sid = pStyle.get(qn('w:val'), '').lower().replace(' ', '')
            for lvl in range(1, 5):
                if any(k in sid for k in (f'heading{lvl}', f'titre{lvl}', f'title{lvl}')):
                    return f'h{lvl}'
                # Certains templates utilisent "1", "2", etc. comme ID de style
                if sid == str(lvl):
                    return f'h{lvl}'
    return 'p'


# =============================================================================
# Collecte des paragraphes (traversée complète du XML)
# =============================================================================

def collect_paragraphs(doc: Document) -> list:
    """
    Retourne la liste plate de tous les éléments <w:p> du corps du document,
    en ordre de lecture, y compris ceux imbriqués dans :
      - des tableaux          : w:body > w:tbl > w:tr > w:tc > w:p
      - des content controls  : w:body > w:sdt > w:sdtContent > w:p
      - des structures mixtes : content controls contenant des tableaux, etc.

    À utiliser à la place de doc.paragraphs, qui ne retourne que les paragraphes
    enfants directs de w:body (paragraphes de premier niveau).

    Returns:
        Liste d'éléments lxml <w:p> dans l'ordre du document. Les indices de
        cette liste sont utilisés comme référence dans toute l'application
        (listbox, highlight, insertion).
    """
    return list(doc.element.body.iter(qn('w:p')))


# =============================================================================
# Prévisualisation HTML
# =============================================================================

def build_html(
    doc: Document,
    highlight_idx: int | None = None,
    flat_paras: list | None = None,
) -> str:
    """
    Génère un fragment HTML à partir des paragraphes du document.

    Chaque paragraphe non vide reçoit un attribut id="para-{i}" où i est son
    index dans flat_paras. Cela permet à l'interface de faire le lien entre
    la listbox (qui affiche les index) et la prévisualisation HTML.

    Le paragraphe highlight_idx reçoit class="highlight" pour être mis en
    évidence visuellement (fond jaune + barre latérale orange).

    Args:
        doc: Document python-docx (utilisé uniquement si flat_paras est None).
        highlight_idx: Index du paragraphe à mettre en évidence, ou None.
        flat_paras: Liste plate des éléments <w:p> produite par collect_paragraphs().
                    Si None, se replie sur [p._p for p in doc.paragraphs] (rétrocompat).

    Returns:
        Chaîne HTML (fragment <body>, sans les balises <html>/<head>).
    """
    paras = flat_paras if flat_paras is not None else [p._p for p in doc.paragraphs]
    parts = []
    for i, p_elem in enumerate(paras):
        text = _para_text(p_elem)
        if not text.strip():
            continue
        tag = _para_html_tag(p_elem)
        cls = ' class="highlight"' if i == highlight_idx else ""
        escaped = _html.escape(text)
        parts.append(f'<{tag} id="para-{i}"{cls}>{escaped}</{tag}>')
    return "\n".join(parts)


def build_html_window(
    doc: Document,
    highlight_idx: int,
    before: int = 4,
    after: int = 25,
    flat_paras: list | None = None,
) -> str:
    """
    Génère un fragment HTML centré sur le paragraphe highlight_idx.

    Affiche jusqu'à `before` paragraphes non vides avant la cible et
    `after` paragraphes après. Un marqueur « — … — » est ajouté si du
    contenu est tronqué au début ou à la fin.

    La cible reçoit class="highlight". Comme la cible est toujours proche
    du haut du fragment, aucun scroll n'est nécessaire côté UI.

    Args:
        doc: Document python-docx (utilisé si flat_paras est None).
        highlight_idx: Index de la cible dans flat_paras.
        before: Nombre max de paragraphes non vides avant la cible.
        after: Nombre max de paragraphes non vides après la cible.
        flat_paras: Liste plate produite par collect_paragraphs().

    Returns:
        Fragment HTML (sans <html>/<head>).
    """
    paras = flat_paras if flat_paras is not None else [p._p for p in doc.paragraphs]
    parts = []

    # ── Paragraphes avant la cible ────────────────────────────────────────────
    before_items: list[tuple[int, object]] = []
    for i in range(highlight_idx - 1, -1, -1):
        if _para_text(paras[i]).strip():
            before_items.append((i, paras[i]))
            if len(before_items) >= before:
                break
    before_items.reverse()

    # Marqueur si le document contient du contenu avant la fenêtre
    n_before_total = sum(1 for i in range(highlight_idx) if _para_text(paras[i]).strip())
    if n_before_total > len(before_items):
        parts.append('<p class="ellipsis">— … —</p>')

    for i, p_elem in before_items:
        tag = _para_html_tag(p_elem)
        escaped = _html.escape(_para_text(p_elem))
        parts.append(f'<{tag} id="para-{i}">{escaped}</{tag}>')

    # ── Paragraphe cible (toujours visible, en haut de la zone) ──────────────
    tag = _para_html_tag(paras[highlight_idx])
    escaped = _html.escape(_para_text(paras[highlight_idx]))
    parts.append(f'<{tag} id="para-{highlight_idx}" class="highlight">{escaped}</{tag}>')

    # ── Paragraphes après la cible ────────────────────────────────────────────
    after_count = 0
    for i in range(highlight_idx + 1, len(paras)):
        if not _para_text(paras[i]).strip():
            continue
        tag = _para_html_tag(paras[i])
        escaped = _html.escape(_para_text(paras[i]))
        parts.append(f'<{tag} id="para-{i}">{escaped}</{tag}>')
        after_count += 1
        if after_count >= after:
            # Marqueur si contenu restant après la fenêtre
            if any(_para_text(paras[j]).strip() for j in range(i + 1, len(paras))):
                parts.append('<p class="ellipsis">— … —</p>')
            break

    return "\n".join(parts)


# =============================================================================
# Gestion des fichiers
# =============================================================================

def open_document(filepath: str) -> Document:
    """
    Ouvre un document .docx en lecture/affichage sans le modifier.

    Le fichier original n'est jamais touché. Les versions modifiées sont
    écrites dans des dossiers de sortie séparés (_track_changes/ et _texte_brut/).
    """
    return Document(filepath)


def save_document(doc: Document, filepath: str) -> None:
    """Sauvegarde le document à son emplacement d'origine (écrase le fichier)."""
    doc.save(filepath)


# =============================================================================
# Recherche et navigation dans les paragraphes
# =============================================================================

def search_paragraphs(
    doc: Document,
    keyword: str,
    flat_paras: list | None = None,
) -> list[tuple[int, str]]:
    """
    Retourne la liste des (index, texte) des paragraphes contenant le mot-clé.

    Args:
        doc: Document python-docx (utilisé uniquement si flat_paras est None).
        keyword: Terme à rechercher (insensible à la casse).
        flat_paras: Liste plate produite par collect_paragraphs().
    """
    paras = flat_paras if flat_paras is not None else [p._p for p in doc.paragraphs]
    keyword_lower = keyword.strip().lower()
    results = []
    for i, p_elem in enumerate(paras):
        text = _para_text(p_elem)
        if keyword_lower and keyword_lower in text.lower():
            results.append((i, text))
    return results


def get_paragraphs_around(
    doc: Document,
    center_idx: int,
    context: int = 4,
    flat_paras: list | None = None,
) -> list[tuple[int, str]]:
    """
    Retourne les paragraphes situés dans une fenêtre de ±context autour
    de center_idx (utilisé pour afficher le contexte d'une occurrence de recherche).

    Args:
        doc: Document python-docx (utilisé uniquement si flat_paras est None).
        center_idx: Index central (paragraphe trouvé par la recherche).
        context: Nombre de paragraphes à afficher de chaque côté.
        flat_paras: Liste plate produite par collect_paragraphs().
    """
    paras = flat_paras if flat_paras is not None else [p._p for p in doc.paragraphs]
    start = max(0, center_idx - context)
    end = min(len(paras), center_idx + context + 1)
    return [(i, _para_text(paras[i])) for i in range(start, end)]


def get_all_paragraphs(
    doc: Document,
    flat_paras: list | None = None,
) -> list[tuple[int, str]]:
    """
    Retourne tous les paragraphes non vides du document avec leur index.

    Args:
        doc: Document python-docx (utilisé uniquement si flat_paras est None).
        flat_paras: Liste plate produite par collect_paragraphs().
    """
    paras = flat_paras if flat_paras is not None else [p._p for p in doc.paragraphs]
    result = []
    for i, p_elem in enumerate(paras):
        text = _para_text(p_elem)
        if text.strip():
            result.append((i, text))
    return result


# =============================================================================
# Insertion de clause en mode révision (Track Changes)
# =============================================================================

def _next_revision_id(doc: Document) -> int:
    """
    Calcule le prochain identifiant de révision disponible dans le document.

    Chaque élément <w:ins> / <w:del> doit porter un w:id unique. On scanne
    le XML brut pour trouver le maximum existant et on retourne max + 1.
    """
    existing_ids = [int(m) for m in re.findall(r'w:id="(\d+)"', doc.element.xml)]
    return (max(existing_ids) + 1) if existing_ids else 1


def _make_tracked_paragraph(
    text: str,
    author: str,
    date_str: str,
    rev_id: int,
    bold: bool = False,
    underline: bool = False,
    font_size: int = 0,
    style_name: str | None = None,
    indent_twips: int = 0,
) -> tuple[OxmlElement, int]:
    """
    Construit un <w:p> dont le contenu est balisé comme insertion Track Changes.

    Structure XML produite :
        <w:p>
          <w:pPr>
            [<w:pStyle w:val="..."/>]       ← si style_name fourni
            [<w:ind w:left="..."/>]         ← si indent_twips > 0
            <w:rPr>
              <w:ins w:id="N" w:author="..." w:date="..."/>
            </w:rPr>
          </w:pPr>
          <w:ins w:id="N+1" w:author="..." w:date="...">
            <w:r>
              [<w:rPr>
                [<w:b/>]                   ← si bold=True
                [<w:u w:val="single"/>]    ← si underline=True
                [<w:sz w:val="N*2"/>]      ← si font_size > 0 (OOXML en demi-points)
                [<w:szCs w:val="N*2"/>]    ← idem pour les scripts complexes
              </w:rPr>]
              <w:t>texte</w:t>
            </w:r>
          </w:ins>
        </w:p>

    Args:
        text: Texte du paragraphe.
        author: Auteur affiché dans la bulle de révision Word.
        date_str: Date ISO 8601 UTC (ex. "2026-03-10T14:32:00Z").
        rev_id: Premier ID de révision disponible (deux IDs consommés).
        bold: Applique le gras au run (sous-titres en mode "Gras").
        underline: Applique le soulignement simple au run (sous-titres en mode "Souligné").
        font_size: Taille de police en points (ex. 12). 0 = hérite du style/document.
                   OOXML utilise des demi-points : 12pt → w:sz val="24".
        style_name: Nom du style Word à appliquer (ex. "Heading 3", "Normal").
                    None = pas de style explicite (hérite du document).
        indent_twips: Indentation gauche en twips pour les puces (0 = aucune).

    Returns:
        Tuple (élément <w:p>, prochain rev_id disponible).
    """
    new_p = OxmlElement("w:p")

    # ── Propriétés de paragraphe ──────────────────────────────────────────────
    pPr = OxmlElement("w:pPr")

    if style_name:
        # Applique un style Word nommé (Heading 3, Normal, Titre 2, etc.)
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_name)
        pPr.append(pStyle)

    if indent_twips > 0:
        # Indentation gauche pour les paragraphes à puce
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent_twips))
        pPr.append(ind)

    # Marque le signe de paragraphe (¶) lui-même comme inséré
    pRpr = OxmlElement("w:rPr")
    ins_pmark = OxmlElement("w:ins")
    ins_pmark.set(qn("w:id"),     str(rev_id))
    ins_pmark.set(qn("w:author"), author)
    ins_pmark.set(qn("w:date"),   date_str)
    pRpr.append(ins_pmark)
    pPr.append(pRpr)
    new_p.append(pPr)

    # ── Run de texte encapsulé dans <w:ins> ───────────────────────────────────
    ins_run = OxmlElement("w:ins")
    ins_run.set(qn("w:id"),     str(rev_id + 1))
    ins_run.set(qn("w:author"), author)
    ins_run.set(qn("w:date"),   date_str)

    new_r = OxmlElement("w:r")
    if bold or underline or font_size:
        # On regroupe toutes les propriétés de caractère dans un seul <w:rPr>
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        if underline:
            # w:u w:val="single" = soulignement simple (le plus courant dans Word)
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
        if font_size:
            # OOXML exprime la taille en demi-points : 12pt → val="24"
            # w:sz = scripts latins ; w:szCs = scripts complexes (arabe, hébreu…)
            half_pts = str(font_size * 2)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), half_pts)
            rPr.append(sz)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), half_pts)
            rPr.append(szCs)
        new_r.append(rPr)

    new_t = OxmlElement("w:t")
    new_t.text = text
    if text.startswith(" ") or text.endswith(" "):
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)

    ins_run.append(new_r)
    new_p.append(ins_run)

    return new_p, rev_id + 2


def insert_clause_after(
    doc: Document,
    para_idx: int,
    subtitle: str,
    text: str,
    author: str,
    subtitle_config: dict | None = None,
    text_style: str | None = None,
    subtitle_font_size: int = 0,
    text_font_size: int = 0,
    flat_paras: list | None = None,
) -> None:
    """
    Insère une clause en mode révision Word (Track Changes) immédiatement
    après le paragraphe désigné par para_idx.

    Utilise flat_paras[para_idx] comme élément ancre si flat_paras est fourni,
    ce qui permet d'insérer après des paragraphes dans des tableaux ou des
    content controls. Sinon, se replie sur doc.paragraphs[para_idx]._p.

    Format du sous-titre contrôlé par subtitle_config :
        {"type": "bold"}                              → texte gras (défaut)
        {"type": "underline"}                         → texte souligné
        {"type": "style", "style": "Heading 3"}       → style Word nommé
        {"type": "puce",  "bullet": "•", "indent": 1} → puce avec indentation

    Args:
        doc: Document python-docx ouvert.
        para_idx: Index dans flat_paras (ou doc.paragraphs si flat_paras=None).
        subtitle: Texte du sous-titre (ignoré si vide).
        text: Corps de la clause.
        author: Nom affiché dans la bulle de révision.
        subtitle_config: Dict décrivant le format du sous-titre (voir ci-dessus).
        text_style: Style Word du corps de clause (ex. "Normal"). None = défaut.
        subtitle_font_size: Taille de police du sous-titre en points (0 = auto).
        text_font_size: Taille de police du corps de clause en points (0 = auto).
        flat_paras: Liste plate produite par collect_paragraphs(). Recommandé.
    """
    # Élément XML ancre — l'insertion se fait via addnext() juste après lui
    if flat_paras is not None:
        anchor = flat_paras[para_idx]
    else:
        anchor = doc.paragraphs[para_idx]._p

    # Résolution du style "auto" : hérite du style du corps de texte voisin.
    # Si l'ancre est un titre (Heading, APU_Heading…), on ne copie PAS son style
    # — ça donnerait une police titre pour le corps de clause. On cherche alors
    # le premier style de corps de texte dans les paragraphes environnants.
    if text_style == "auto":
        anchor_style = get_para_style_name(anchor)
        paras = flat_paras if flat_paras is not None else [p._p for p in doc.paragraphs]
        if anchor_style and _is_heading_style(anchor_style):
            resolved_text_style = get_body_style_near(paras, para_idx) or None
            # Taille de police : lire depuis le paragraphe corps voisin si non forcée
            if text_font_size == 0:
                body_para = _find_body_para_near(paras, para_idx)
                effective_text_font = get_para_run_font_size(body_para) if body_para is not None else 0
            else:
                effective_text_font = text_font_size
        else:
            resolved_text_style = anchor_style or None
            # Taille de police : lire depuis l'ancre si non forcée
            effective_text_font = text_font_size if text_font_size else get_para_run_font_size(anchor)
    else:
        resolved_text_style = _resolve_style_id(doc, text_style) if text_style else None
        effective_text_font = text_font_size  # style explicite : respecter la valeur de l'utilisateur

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    rev_id = _next_revision_id(doc)

    cfg = subtitle_config or {}
    sub_type = cfg.get("type", "bold")

    # Construire la liste des paragraphes à insérer (sous-titre + corps)
    items = []

    if subtitle and subtitle.strip():
        if sub_type == "style":
            items.append({
                "text": subtitle.strip(),
                "bold": False, "underline": False,
                "font_size": subtitle_font_size,
                "style_name": _resolve_heading_style_id(doc, cfg.get("style", "Heading 3")),
                "indent_twips": 0,
            })
        elif sub_type == "puce":
            bullet = cfg.get("bullet", "•")
            level = int(cfg.get("indent", 1))
            items.append({
                "text": f"{bullet}\t{subtitle.strip()}",
                "bold": False, "underline": False,
                "font_size": subtitle_font_size,
                "style_name": None,
                "indent_twips": INDENT_LEVELS.get(level, 720),
            })
        elif sub_type == "underline":
            items.append({
                "text": subtitle.strip(),
                "bold": False, "underline": True,
                "font_size": subtitle_font_size,
                "style_name": None,
                "indent_twips": 0,
            })
        else:  # "bold" (défaut)
            items.append({
                "text": subtitle.strip(),
                "bold": True, "underline": False,
                "font_size": subtitle_font_size,
                "style_name": None,
                "indent_twips": 0,
            })

    items.append({
        "text": text.strip(),
        "bold": False, "underline": False,
        "font_size": effective_text_font,
        "style_name": resolved_text_style,
        "indent_twips": 0,
    })

    # Insertion en ordre inversé pour que addnext() produise l'ordre final correct.
    # addnext(X) insère X immédiatement après l'ancre. En insérant dans l'ordre
    # [corps, sous-titre] (inversé), le résultat dans le document est [sous-titre, corps].
    for item in reversed(items):
        new_p, rev_id = _make_tracked_paragraph(
            item["text"], author, date_str, rev_id,
            bold=item["bold"],
            underline=item["underline"],
            font_size=item["font_size"],
            style_name=item["style_name"],
            indent_twips=item["indent_twips"],
        )
        anchor.addnext(new_p)


# =============================================================================
# Insertion de clause en texte brut (sans Track Changes)
# =============================================================================

def _make_plain_paragraph(
    text: str,
    bold: bool = False,
    underline: bool = False,
    font_size: int = 0,
    style_name: str | None = None,
    indent_twips: int = 0,
) -> OxmlElement:
    """
    Construit un <w:p> avec le texte inséré directement, sans balisage
    Track Changes (<w:ins>). Utilisé pour la version "texte brut".

    Structure XML produite :
        <w:p>
          <w:pPr>
            [<w:pStyle w:val="..."/>]       ← si style_name fourni
            [<w:ind w:left="..."/>]         ← si indent_twips > 0
          </w:pPr>
          <w:r>
            [<w:rPr>
              [<w:b/>]
              [<w:u w:val="single"/>]
              [<w:sz/><w:szCs/>]
            </w:rPr>]
            <w:t>texte</w:t>
          </w:r>
        </w:p>

    Args:
        text: Texte du paragraphe.
        bold: Applique le gras.
        underline: Applique le soulignement simple.
        font_size: Taille en points (0 = hérite du style).
        style_name: Nom du style Word (None = hérite).
        indent_twips: Indentation gauche en twips (0 = aucune).

    Returns:
        Élément <w:p> prêt à être inséré via addnext().
    """
    new_p = OxmlElement("w:p")

    # ── Propriétés de paragraphe ──────────────────────────────────────────────
    pPr = OxmlElement("w:pPr")
    if style_name:
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_name)
        pPr.append(pStyle)
    if indent_twips > 0:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent_twips))
        pPr.append(ind)
    new_p.append(pPr)

    # ── Run de texte direct (pas de <w:ins>) ──────────────────────────────────
    new_r = OxmlElement("w:r")
    if bold or underline or font_size:
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
        if font_size:
            half_pts = str(font_size * 2)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), half_pts)
            rPr.append(sz)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), half_pts)
            rPr.append(szCs)
        new_r.append(rPr)

    new_t = OxmlElement("w:t")
    new_t.text = text
    if text.startswith(" ") or text.endswith(" "):
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)

    return new_p


def insert_clause_plain_after(
    doc: Document,
    para_idx: int,
    subtitle: str,
    text: str,
    subtitle_config: dict | None = None,
    text_style: str | None = None,
    subtitle_font_size: int = 0,
    text_font_size: int = 0,
    flat_paras: list | None = None,
) -> None:
    """
    Insère une clause en texte brut (sans Track Changes) immédiatement après
    le paragraphe désigné par para_idx. Le résultat est un document Word
    classique, sans bulle de révision.

    Même logique de construction que insert_clause_after, mais les paragraphes
    sont créés avec _make_plain_paragraph au lieu de _make_tracked_paragraph :
    pas de <w:ins>, pas d'auteur, pas de date de révision.

    Args:
        doc: Document python-docx ouvert (copie fraîche de l'original).
        para_idx: Index dans flat_paras (ou doc.paragraphs si flat_paras=None).
        subtitle: Texte du sous-titre (ignoré si vide).
        text: Corps de la clause.
        subtitle_config: Dict décrivant le format du sous-titre.
        text_style: Style Word du corps de clause (ex. "Normal"). None = défaut.
        subtitle_font_size: Taille de police du sous-titre en points (0 = auto).
        text_font_size: Taille de police du corps de clause en points (0 = auto).
        flat_paras: Liste plate produite par collect_paragraphs(). Recommandé.
    """
    if flat_paras is not None:
        anchor = flat_paras[para_idx]
    else:
        anchor = doc.paragraphs[para_idx]._p

    # Résolution du style "auto" : hérite du style du corps de texte voisin.
    # Si l'ancre est un titre (Heading, APU_Heading…), on ne copie PAS son style
    # — ça donnerait une police titre pour le corps de clause. On cherche alors
    # le premier style de corps de texte dans les paragraphes environnants.
    if text_style == "auto":
        anchor_style = get_para_style_name(anchor)
        paras = flat_paras if flat_paras is not None else [p._p for p in doc.paragraphs]
        if anchor_style and _is_heading_style(anchor_style):
            resolved_text_style = get_body_style_near(paras, para_idx) or None
            if text_font_size == 0:
                body_para = _find_body_para_near(paras, para_idx)
                effective_text_font = get_para_run_font_size(body_para) if body_para is not None else 0
            else:
                effective_text_font = text_font_size
        else:
            resolved_text_style = anchor_style or None
            effective_text_font = text_font_size if text_font_size else get_para_run_font_size(anchor)
    else:
        resolved_text_style = _resolve_style_id(doc, text_style) if text_style else None
        effective_text_font = text_font_size

    cfg = subtitle_config or {}
    sub_type = cfg.get("type", "bold")
    items = []

    if subtitle and subtitle.strip():
        if sub_type == "style":
            items.append({
                "text": subtitle.strip(),
                "bold": False, "underline": False,
                "font_size": subtitle_font_size,
                "style_name": _resolve_heading_style_id(doc, cfg.get("style", "Heading 3")),
                "indent_twips": 0,
            })
        elif sub_type == "puce":
            bullet = cfg.get("bullet", "•")
            level = int(cfg.get("indent", 1))
            items.append({
                "text": f"{bullet}\t{subtitle.strip()}",
                "bold": False, "underline": False,
                "font_size": subtitle_font_size,
                "style_name": None,
                "indent_twips": INDENT_LEVELS.get(level, 720),
            })
        elif sub_type == "underline":
            items.append({
                "text": subtitle.strip(),
                "bold": False, "underline": True,
                "font_size": subtitle_font_size,
                "style_name": None,
                "indent_twips": 0,
            })
        else:  # "bold" (défaut)
            items.append({
                "text": subtitle.strip(),
                "bold": True, "underline": False,
                "font_size": subtitle_font_size,
                "style_name": None,
                "indent_twips": 0,
            })

    items.append({
        "text": text.strip(),
        "bold": False, "underline": False,
        "font_size": effective_text_font,
        "style_name": resolved_text_style,
        "indent_twips": 0,
    })

    for item in reversed(items):
        new_p = _make_plain_paragraph(
            item["text"],
            bold=item["bold"],
            underline=item["underline"],
            font_size=item["font_size"],
            style_name=item["style_name"],
            indent_twips=item["indent_twips"],
        )
        anchor.addnext(new_p)


# =============================================================================
# Mise à jour des dates de publication
#
# Deux patterns sont détectés automatiquement dans chaque document :
#   1. Corps     : paragraphe contenant "Date de publication"
#                  ex. "Date de publication : 01/01/2026"
#   2. Footer    : paragraphe contenant "mise à jour le"
#                  ex. "Dernière mise à jour le 28/11/2025"
#
# La date trouvée (format JJ/MM/AAAA) est remplacée par la date du jour.
# En mode Track Changes : <w:del> (ancienne date) + <w:ins> (nouvelle date),
# les propriétés de caractère du run original (w:rPr) sont conservées.
# En mode texte brut : remplacement direct dans <w:t>.
# =============================================================================

_DATE_RE = re.compile(r'\d{2}/\d{2}/\d{4}')


def _collect_visible_runs(el, result: list) -> None:
    """
    Collecte récursivement les runs visibles dans el, en ignorant <w:del>.

    Appelé par _visible_runs pour traverser les structures imbriquées :
    <w:hyperlink>, <w:ins>, <w:sdt>, <w:sdtContent>, etc.
    """
    for child in el:
        if child.tag == qn('w:del'):
            continue  # texte supprimé — invisible dans le document
        if child.tag == qn('w:r'):
            t = child.find(qn('w:t'))
            result.append((child, t.text or '' if t is not None else ''))
        else:
            _collect_visible_runs(child, result)


def _visible_runs(p_elem) -> list[tuple]:
    """
    Retourne la liste des (run_el, texte) pour les runs visibles du paragraphe.

    Récurse dans tous les éléments enfants sauf <w:del> (texte supprimé).
    Couvre les runs directs, dans <w:ins>, <w:hyperlink>, <w:sdt>, etc.
    Les runs dans <w:del> sont exclus car ils représentent du texte supprimé.
    """
    result = []
    _collect_visible_runs(p_elem, result)
    return result


def _find_date_run(p_elem):
    """
    Cherche un pattern JJ/MM/AAAA dans les runs visibles du paragraphe.

    La recherche est effectuée dans le texte combiné de tous les runs visibles
    (y compris les runs dans <w:hyperlink>, <w:ins>, <w:sdt>…). La date doit
    être entièrement contenue dans un seul run pour permettre le remplacement.
    Si elle s'étale sur plusieurs runs, retourne None.

    Returns:
        (run_el, start_in_run, end_in_run, old_date_str) ou None.
    """
    runs = _visible_runs(p_elem)
    combined = ''.join(text for _, text in runs)
    m = _DATE_RE.search(combined)
    if not m:
        return None
    pos = 0
    for run_el, text in runs:
        run_end = pos + len(text)
        if pos <= m.start() and m.end() <= run_end:
            return run_el, m.start() - pos, m.end() - pos, m.group()
        pos = run_end
    return None  # date répartie sur plusieurs runs — non supporté


def _replace_date_in_run_plain(run_el, old_date: str, new_date: str) -> None:
    """Remplace old_date par new_date directement dans le <w:t> du run."""
    t_el = run_el.find(qn('w:t'))
    if t_el is not None and t_el.text:
        t_el.text = t_el.text.replace(old_date, new_date, 1)


def _replace_date_in_run_tracked(
    run_el, old_date: str, new_date: str,
    author: str, date_utc: str, rev_id: int,
) -> int:
    """
    Remplace old_date par new_date dans run_el avec balisage Track Changes.

    Décompose le run en [prefix?] <w:del>ancienne</w:del>
    <w:ins>nouvelle</w:ins> [suffix?], en conservant les propriétés de
    caractère (w:rPr) du run original dans tous les nouveaux éléments.

    Returns:
        Prochain rev_id disponible (consomme 2 IDs : un pour del, un pour ins).
    """
    parent    = run_el.getparent()
    t_el      = run_el.find(qn('w:t'))
    if t_el is None:
        return rev_id

    full_text = t_el.text or ''
    idx = full_text.find(old_date)
    if idx < 0:
        return rev_id

    prefix   = full_text[:idx]
    suffix   = full_text[idx + len(old_date):]
    rPr_orig = run_el.find(qn('w:rPr'))

    def _make_run(text: str) -> OxmlElement:
        r = OxmlElement('w:r')
        if rPr_orig is not None:
            r.append(deepcopy(rPr_orig))
        t = OxmlElement('w:t')
        t.text = text
        if text and (text[0] == ' ' or text[-1] == ' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        return r

    def _make_del_run(text: str) -> OxmlElement:
        r = OxmlElement('w:r')
        if rPr_orig is not None:
            r.append(deepcopy(rPr_orig))
        dt = OxmlElement('w:delText')
        dt.text = text
        r.append(dt)
        return r

    insert_pos = list(parent).index(run_el)
    elements   = []

    if prefix:
        elements.append(_make_run(prefix))

    del_el = OxmlElement('w:del')
    del_el.set(qn('w:id'),     str(rev_id))
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:date'),   date_utc)
    del_el.append(_make_del_run(old_date))
    elements.append(del_el)

    ins_el = OxmlElement('w:ins')
    ins_el.set(qn('w:id'),     str(rev_id + 1))
    ins_el.set(qn('w:author'), author)
    ins_el.set(qn('w:date'),   date_utc)
    ins_el.append(_make_run(new_date))
    elements.append(ins_el)

    if suffix:
        elements.append(_make_run(suffix))

    parent.remove(run_el)
    for i, el in enumerate(elements):
        parent.insert(insert_pos + i, el)

    return rev_id + 2


def _update_dates_in_doc(
    doc: Document,
    author: str | None,
    flat_paras: list | None,
    tracked: bool,
) -> dict:
    """
    Logique commune de mise à jour des dates (corps + footer).

    Cherche :
      - Corps  : premier paragraphe contenant "Date de publication"
      - Footer : premier paragraphe de footer contenant "mise à jour le"
        (insensible à la casse, couvre "Dernière mise à jour le")

    Si la date trouvée est déjà celle d'aujourd'hui, aucune modification n'est
    effectuée (évite les révisions inutiles).

    Plusieurs sections peuvent partager le même footer XML (lié au précédent) :
    on déduplique par identité de l'élément <w:ftr> pour ne le traiter qu'une
    seule fois.

    Returns:
        {"body": bool, "footer": bool} — True si le pattern a été trouvé
        (indépendamment du fait que la date ait été modifiée ou non).
    """
    today    = datetime.now().strftime("%d/%m/%Y")
    date_utc = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    rev_id   = _next_revision_id(doc) if tracked else 0
    result   = {"body": False, "footer": False}

    paras = flat_paras if flat_paras is not None else [p._p for p in doc.paragraphs]

    # ── Corps : "Date de publication" ─────────────────────────────────────────
    for p_el in paras:
        if "Date de publication" in _para_text(p_el):
            found = _find_date_run(p_el)
            if found:
                run_el, _, _, old_date = found
                if old_date != today:
                    if tracked:
                        rev_id = _replace_date_in_run_tracked(
                            run_el, old_date, today, author, date_utc, rev_id)
                    else:
                        _replace_date_in_run_plain(run_el, old_date, today)
            result["body"] = True
            break  # première occurrence seulement

    # ── Footer : "mise à jour le" ─────────────────────────────────────────────
    seen_ftr = set()
    for section in doc.sections:
        try:
            ftr_el = section.footer._element
        except Exception:
            continue
        if id(ftr_el) in seen_ftr:
            continue
        seen_ftr.add(id(ftr_el))

        for p_el in ftr_el.iter(qn('w:p')):
            if "mise à jour le" in _para_text(p_el).lower():
                found = _find_date_run(p_el)
                if found:
                    run_el, _, _, old_date = found
                    if old_date != today:
                        if tracked:
                            rev_id = _replace_date_in_run_tracked(
                                run_el, old_date, today, author, date_utc, rev_id)
                        else:
                            _replace_date_in_run_plain(run_el, old_date, today)
                result["footer"] = True
                break

        if result["footer"]:
            break

    return result


def update_dates(doc: Document, author: str, flat_paras: list | None = None) -> dict:
    """
    Met à jour les dates de publication avec Track Changes (<w:del> + <w:ins>).

    Args:
        doc: Document python-docx (modifié en place).
        author: Auteur affiché dans les bulles de révision.
        flat_paras: Liste plate produite par collect_paragraphs(). Recommandé.

    Returns:
        {"body": bool, "footer": bool} — True si le pattern a été trouvé.
    """
    return _update_dates_in_doc(doc, author, flat_paras, tracked=True)


def update_dates_plain(doc: Document, flat_paras: list | None = None) -> dict:
    """
    Met à jour les dates de publication en texte direct (sans Track Changes).

    Args:
        doc: Document python-docx (modifié en place).
        flat_paras: Liste plate produite par collect_paragraphs(). Recommandé.

    Returns:
        {"body": bool, "footer": bool} — True si le pattern a été trouvé.
    """
    return _update_dates_in_doc(doc, None, flat_paras, tracked=False)